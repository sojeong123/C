from docx import Document

def remove_double_slash(file_path):
    doc = Document(file_path)

    for paragraph in doc.paragraphs:
        updated_text = paragraph.text.replace("//", "")
        paragraph.clear()
        paragraph.add_run(updated_text)

    doc.save(file_path)

if __name__ == "__main__":
    file_path = r"C:\Users\Administrator\Desktop\3\midterm.docx"
    remove_double_slash(file_path)
